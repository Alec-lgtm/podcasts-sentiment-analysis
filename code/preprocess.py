import os
rt os
from pathlib import Path
from typing import List, Tuple, Dict, Set
import re
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize
import pickle

def setup_nltk():
    """Download required NLTK data if not already present."""
    try:
        nltk.data.find('corpora/stopwords')
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('stopwords', quiet=True)
        nltk.download('punkt', quiet=True)

def load_nrc_lexicon(negative_words: Set[str], neutral_words: Set[str]) -> Dict[str, int]:
    """Create lexicon dictionary from the provided sets."""
    lexicon = {}
    for word in negative_words:
        lexicon[word] = 1  # negative sentiment
    for word in neutral_words:
        lexicon[word] = 0  # neutral/positive sentiment
    return lexicon

def tokenize_sentence(sentence: str, lexicon: Dict[str, int]) -> List[str]:
    """Convert a sentence into a list of tokens."""
    stop_words = set(stopwords.words('english'))

    # Clean the text
    sentence = re.sub(r'\W+', ' ', sentence)

    # Convert to lowercase and split into tokens
    tokens = sentence.lower().split()

    # Filter stop words, single characters, and keep only words in our lexicon
    return [word for word in tokens if word not in stop_words and len(word) > 1 and word in lexicon]

def classify_sentence(tokens: List[str], lexicon: Dict[str, int]) -> int:
    """Classify a sentence based on the lexicon."""
    if not tokens:
        return 0

    # Count negative words
    negative_count = sum(1 for token in tokens if lexicon.get(token, 0) == 1)

    # If more than 20% of the words are negative, classify as negative
    return 1 if negative_count / len(tokens) > 0.2 else 0

def process_documents_for_naive_bayes(directory: str, lexicon: Dict[str, int]) -> List[Tuple[List[str], int]]:
    """
    Process documents and return sentence-level data for Naive Bayes.
    Each sentence becomes a separate training example.
    """
    setup_nltk()
    training_data = []
    directory_path = Path(directory)

    if not directory_path.exists():
        raise FileNotFoundError(f"Directory not found: {directory}")

    for filepath in directory_path.glob("*.docx"):
        try:
            # Process document
            doc = Document(filepath)
            full_text = " ".join(paragraph.text for paragraph in doc.paragraphs)

            # Split into sentences
            sentences = sent_tokenize(full_text)

            # Process each sentence
            for sentence in sentences:
                tokens = tokenize_sentence(sentence, lexicon)
                if tokens and len(tokens) >= 3:  # Only add if we have enough tokens
                    class_label = classify_sentence(tokens, lexicon)
                    training_data.append((tokens, class_label))

            print(f"Processed {filepath.name}")

        except Exception as e:
            print(f"Error processing {filepath.name}: {str(e)}")

    return training_data

def save_training_data(training_data: List[Tuple[List[str], int]], filename: str):
    """Save the training data to a pickle file."""
    with open(filename, 'wb') as f:
        pickle.dump(training_data, f)

def analyze_training_data(training_data: List[Tuple[List[str], int]]):
    """Print analysis of the processed training data."""
    class_counts = {}
    sentence_lengths = []

    for tokens, label in training_data:
        class_counts[label] = class_counts.get(label, 0) + 1
        sentence_lengths.append(len(tokens))

    print("\nTraining Data Analysis:")
    print(f"Total sentences: {len(training_data)}")
    print("\nSentences per class:")
    for label, count in class_counts.items():
        print(f"Class {label}: {count} sentences")

    if sentence_lengths:
        avg_length = sum(sentence_lengths) / len(sentence_lengths)
        print(f"\nAverage tokens per sentence: {avg_length:.1f}")
        print(f"Min tokens per sentence: {min(sentence_lengths)}")
        print(f"Max tokens per sentence: {max(sentence_lengths)}")

def main():
    # Create sets from the provided lexicons
    negative_words = {word.strip() for word, value in [line.strip().split('\t') for line in negative_content.split('\n')] if value == '1'}
    neutral_words = {word.strip() for word, value in [line.strip().split('\t') for line in neutral_content.split('\n')] if value == '0'}

    # Create lexicon dictionary
    lexicon = load_nrc_lexicon(negative_words, neutral_words)

    directory = "../data/vox_podcasts/2024"
    output_file = "training_data.pkl"

    try:
        # Process documents
        training_data = process_documents_for_naive_bayes(directory, lexicon)

        # Save the training data
        save_training_data(training_data, output_file)
        print(f"\nSaved training data to {output_file}")

        # Analyze the training data
        analyze_training_data(training_data)

    except Exception as e:
        print(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
from pathlib import Path
from typing import List, Tuple, Dict, Set
import re
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize
import pickle

def setup_nltk():
    """Download required NLTK data if not already present."""
    try:
        nltk.data.find('corpora/stopwords')
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('stopwords', quiet=True)
        nltk.download('punkt', quiet=True)

def load_nrc_lexicon(negative_words: Set[str], neutral_words: Set[str]) -> Dict[str, int]:
    """Create lexicon dictionary from the provided sets."""
    lexicon = {}
    for word in negative_words:
        lexicon[word] = 1  # negative sentiment
    for word in neutral_words:
        lexicon[word] = 0  # neutral/positive sentiment
    return lexicon

def tokenize_sentence(sentence: str, lexicon: Dict[str, int]) -> List[str]:
    """Convert a sentence into a list of tokens."""
    stop_words = set(stopwords.words('english'))

    # Clean the text
    sentence = re.sub(r'\W+', ' ', sentence)

    # Convert to lowercase and split into tokens
    tokens = sentence.lower().split()

    # Filter stop words, single characters, and keep only words in our lexicon
    return [word for word in tokens if word not in stop_words and len(word) > 1 and word in lexicon]

def classify_sentence(tokens: List[str], lexicon: Dict[str, int]) -> int:
    """Classify a sentence based on the lexicon."""
    if not tokens:
        return 0

    # Count negative words
    negative_count = sum(1 for token in tokens if lexicon.get(token, 0) == 1)

    # If more than 20% of the words are negative, classify as negative
    return 1 if negative_count / len(tokens) > 0.2 else 0

def process_documents_for_naive_bayes(directory: str, lexicon: Dict[str, int]) -> List[Tuple[List[str], int]]:
    """
    Process documents and return sentence-level data for Naive Bayes.
    Each sentence becomes a separate training example.
    """
    setup_nltk()
    training_data = []
    directory_path = Path(directory)

    if not directory_path.exists():
        raise FileNotFoundError(f"Directory not found: {directory}")

    for filepath in directory_path.glob("*.docx"):
        try:
            # Process document
            doc = Document(filepath)
            full_text = " ".join(paragraph.text for paragraph in doc.paragraphs)

            # Split into sentences
            sentences = sent_tokenize(full_text)

            # Process each sentence
            for sentence in sentences:
                tokens = tokenize_sentence(sentence, lexicon)
                if tokens and len(tokens) >= 3:  # Only add if we have enough tokens
                    class_label = classify_sentence(tokens, lexicon)
                    training_data.append((tokens, class_label))

            print(f"Processed {filepath.name}")

        except Exception as e:
            print(f"Error processing {filepath.name}: {str(e)}")

    return training_data

def save_training_data(training_data: List[Tuple[List[str], int]], filename: str):
    """Save the training data to a pickle file."""
    with open(filename, 'wb') as f:
        pickle.dump(training_data, f)

def analyze_training_data(training_data: List[Tuple[List[str], int]]):
    """Print analysis of the processed training data."""
    class_counts = {}
    sentence_lengths = []

    for tokens, label in training_data:
        class_counts[label] = class_counts.get(label, 0) + 1
        sentence_lengths.append(len(tokens))

    print("\nTraining Data Analysis:")
    print(f"Total sentences: {len(training_data)}")
    print("\nSentences per class:")
    for label, count in class_counts.items():
        print(f"Class {label}: {count} sentences")

    if sentence_lengths:
        avg_length = sum(sentence_lengths) / len(sentence_lengths)
        print(f"\nAverage tokens per sentence: {avg_length:.1f}")
        print(f"Min tokens per sentence: {min(sentence_lengths)}")
        print(f"Max tokens per sentence: {max(sentence_lengths)}")

def main():
    # Create sets from the provided lexicons
    negative_words = {word.strip() for word, value in [line.strip().split('\t') for line in negative_content.split('\n')] if value == '1'}
    neutral_words = {word.strip() for word, value in [line.strip().split('\t') for line in neutral_content.split('\n')] if value == '0'}

    # Create lexicon dictionary
    lexicon = load_nrc_lexicon(negative_words, neutral_words)

    directory = "../data/vox_podcasts"
    output_file = "training_data.pkl"

    try:
        # Process documents
        training_data = process_documents_for_naive_bayes(directory, lexicon)

        # Save the training data
        save_training_data(training_data, output_file)
        print(f"\nSaved training data to {output_file}")

        # Analyze the training data
        analyze_training_data(training_data)

    except Exception as e:
        print(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
